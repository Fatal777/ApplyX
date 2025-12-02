"""
Resume Customization Service
==============================
AI-powered resume customization to match job descriptions.

Features:
- AI analysis of job description
- Skills gap identification
- Resume content optimization
- PDF and DOCX export
- Version tracking per job application

PDF Generation Strategy:
- Using PyMuPDF (fitz) which is already in requirements.txt
- Supports professional formatting with sections
- Fast and lightweight
"""

from __future__ import annotations

import io
import json
import logging
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from sqlalchemy.orm import Session

try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.models.application import JobApplication, CustomizedResume
from app.models.resume import Resume
from app.services.credits_service import CreditsService

logger = logging.getLogger(__name__)


class ResumeCustomizationService:
    """
    Service for AI-powered resume customization.
    
    Workflow:
    1. Analyze job description
    2. Match against user's resume
    3. Generate optimized content
    4. Create PDF/DOCX output
    5. Track as new version
    """
    
    def __init__(self, db: Session):
        self.db = db
        self.credits_service = CreditsService(db)
    
    async def customize_resume_for_job(
        self,
        user_id: int,
        application_id: int,
        resume_id: int,
        job_description: str,
        job_title: str,
        company_name: str,
        use_ai: bool = True,
    ) -> Dict[str, Any]:
        """
        Create a customized resume version for a specific job.
        
        Args:
            user_id: User requesting customization
            application_id: Job application ID
            resume_id: Base resume ID
            job_description: Full job description text
            job_title: Target job title
            company_name: Target company
            use_ai: Whether to use AI optimization (costs credit)
        
        Returns:
            Dict with customization result and new version info
        """
        # Check credits if using AI
        if use_ai:
            credit_check = self.credits_service.get_available_credits(user_id)
            if not credit_check["can_customize"]:
                return {
                    "success": False,
                    "error": "No credits available",
                    "credits": credit_check,
                    "upgrade_message": "Upgrade your plan for more customizations",
                }
        
        # Get the original resume
        resume = self.db.query(Resume).filter(Resume.id == resume_id).first()
        if not resume:
            return {"success": False, "error": "Resume not found"}
        
        # Get the application
        application = self.db.query(JobApplication).filter(
            JobApplication.id == application_id,
            JobApplication.user_id == user_id,
        ).first()
        
        if not application:
            return {"success": False, "error": "Application not found"}
        
        # Analyze job description
        analysis = self._analyze_job_description(job_description)
        
        # Match with resume
        match_result = self._match_resume_to_job(resume, analysis)
        
        # Generate optimized content
        if use_ai:
            optimized_content = self._generate_ai_optimized_content(
                resume, analysis, match_result, job_title, company_name
            )
            
            # Use credit
            self.credits_service.use_credit(
                user_id,
                action="resume_customization",
                application_id=application_id,
                description=f"Customized for {job_title} at {company_name}",
            )
        else:
            # Basic optimization without AI
            optimized_content = self._generate_basic_optimization(
                resume, analysis, match_result
            )
        
        # Create customized resume version
        version_number = self._get_next_version_number(resume_id)
        
        customized = CustomizedResume(
            resume_id=resume_id,
            application_id=application_id,
            user_id=user_id,
            version_number=version_number,
            original_content=resume.parsed_data,
            customized_content=optimized_content,
            changes_made=match_result.get("changes_made", []),
            target_job_title=job_title,
            target_company=company_name,
        )
        
        self.db.add(customized)
        self.db.commit()
        self.db.refresh(customized)
        
        return {
            "success": True,
            "customized_resume_id": customized.id,
            "version_number": version_number,
            "analysis": analysis,
            "match_result": match_result,
            "changes_made": match_result.get("changes_made", []),
            "improvement_suggestions": match_result.get("suggestions", []),
        }
    
    def _analyze_job_description(self, job_description: str) -> Dict[str, Any]:
        """Extract key requirements from job description."""
        text = job_description.lower()
        
        # Extract required skills
        skill_patterns = [
            r"(?:required|must have|essential).*?(?:skills?|experience).*?[:]\s*(.*?)(?:\.|$)",
            r"(?:looking for|seeking).*?(?:with|who has)\s*(.*?)(?:\.|$)",
        ]
        
        required_skills = []
        for pattern in skill_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                skills = [s.strip() for s in re.split(r'[,;]', match) if s.strip()]
                required_skills.extend(skills)
        
        # Common tech skills extraction
        tech_keywords = [
            "python", "javascript", "typescript", "react", "node.js", "nodejs",
            "java", "c++", "go", "rust", "sql", "mongodb", "postgresql", "mysql",
            "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "git",
            "machine learning", "deep learning", "ai", "data science",
            "agile", "scrum", "ci/cd", "devops", "microservices", "rest api",
            "html", "css", "vue", "angular", "django", "flask", "fastapi",
            "tensorflow", "pytorch", "pandas", "numpy", "scikit-learn",
        ]
        
        found_skills = [skill for skill in tech_keywords if skill in text]
        required_skills.extend(found_skills)
        required_skills = list(set(required_skills))
        
        # Extract experience requirements
        exp_patterns = [
            r"(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*experience",
            r"experience.*?(\d+)\+?\s*(?:years?|yrs?)",
        ]
        
        experience_years = None
        for pattern in exp_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                experience_years = int(match.group(1))
                break
        
        # Extract education requirements
        education_keywords = ["bachelor", "master", "phd", "b.s.", "m.s.", "b.tech", "m.tech"]
        education_required = any(edu in text for edu in education_keywords)
        
        # Extract job type
        job_types = {
            "remote": "remote" in text,
            "hybrid": "hybrid" in text,
            "onsite": any(w in text for w in ["on-site", "onsite", "in-office"]),
            "full-time": "full-time" in text or "full time" in text,
            "part-time": "part-time" in text or "part time" in text,
            "contract": "contract" in text,
        }
        
        # Extract key responsibilities
        responsibility_patterns = [
            r"responsibilities?.*?[:]\s*(.*?)(?:requirements?|qualifications?|$)",
            r"what you'll do.*?[:]\s*(.*?)(?:requirements?|qualifications?|$)",
            r"role.*?includes?.*?[:]\s*(.*?)(?:requirements?|$)",
        ]
        
        responsibilities = []
        for pattern in responsibility_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                items = [r.strip() for r in re.split(r'[\n•-]', match) if len(r.strip()) > 10]
                responsibilities.extend(items[:5])
        
        return {
            "required_skills": required_skills[:15],  # Top 15 skills
            "experience_years": experience_years,
            "education_required": education_required,
            "job_types": job_types,
            "responsibilities": responsibilities[:8],
            "raw_text_length": len(job_description),
        }
    
    def _match_resume_to_job(
        self, 
        resume: Resume, 
        job_analysis: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Match resume content against job requirements."""
        resume_skills = resume.skills or []
        resume_skills_lower = [s.lower() for s in resume_skills]
        
        required_skills = job_analysis.get("required_skills", [])
        
        # Skills matching
        matched_skills = []
        missing_skills = []
        
        for skill in required_skills:
            skill_lower = skill.lower()
            if any(skill_lower in rs or rs in skill_lower for rs in resume_skills_lower):
                matched_skills.append(skill)
            else:
                missing_skills.append(skill)
        
        # Calculate match score
        if required_skills:
            skill_match_score = len(matched_skills) / len(required_skills) * 100
        else:
            skill_match_score = 70  # Default if no skills extracted
        
        # Experience matching
        experience_match = True
        experience_note = ""
        jd_exp = job_analysis.get("experience_years")
        resume_exp = getattr(resume, "experience_years", None)
        
        if jd_exp and resume_exp:
            if resume_exp < jd_exp:
                experience_match = False
                experience_note = f"Resume shows {resume_exp}yrs, job needs {jd_exp}yrs"
        
        # Generate improvement suggestions
        suggestions = []
        changes_made = []
        
        if missing_skills:
            suggestions.append({
                "type": "skills",
                "priority": "high",
                "message": f"Add these skills if you have them: {', '.join(missing_skills[:5])}",
            })
            
            # Auto-add skills user might have
            for skill in missing_skills[:3]:
                changes_made.append({
                    "section": "skills",
                    "action": "add",
                    "value": skill,
                    "reason": "Required by job description",
                })
        
        if not experience_match:
            suggestions.append({
                "type": "experience",
                "priority": "medium",
                "message": experience_note,
            })
        
        # Suggest keyword optimization
        suggestions.append({
            "type": "keywords",
            "priority": "high",
            "message": "Resume should include these keywords prominently",
            "keywords": matched_skills[:5],
        })
        
        return {
            "match_score": round(skill_match_score, 1),
            "matched_skills": matched_skills,
            "missing_skills": missing_skills,
            "experience_match": experience_match,
            "suggestions": suggestions,
            "changes_made": changes_made,
        }
    
    def _generate_ai_optimized_content(
        self,
        resume: Resume,
        job_analysis: Dict[str, Any],
        match_result: Dict[str, Any],
        job_title: str,
        company_name: str,
    ) -> Dict[str, Any]:
        """Generate AI-optimized resume content."""
        original = resume.parsed_data or {}
        optimized = original.copy()
        
        # Optimize summary/objective
        if "summary" in optimized:
            original_summary = optimized["summary"]
            # Add job-specific keywords
            keywords = match_result.get("matched_skills", [])[:3]
            if keywords:
                keyword_text = ", ".join(keywords)
                optimized["summary"] = f"{original_summary} Experienced in {keyword_text}."
        
        # Optimize skills section
        skills = list(optimized.get("skills", []))
        missing = match_result.get("missing_skills", [])
        
        # Add relevant missing skills (user should verify they have these)
        for skill in missing[:5]:
            if skill not in skills:
                skills.insert(0, skill)  # Add to top
        
        optimized["skills"] = skills
        
        # Add job-targeted metadata
        optimized["_optimization"] = {
            "target_job_title": job_title,
            "target_company": company_name,
            "optimized_at": datetime.utcnow().isoformat(),
            "keywords_added": missing[:5],
            "original_match_score": match_result.get("match_score", 0),
        }
        
        return optimized
    
    def _generate_basic_optimization(
        self,
        resume: Resume,
        job_analysis: Dict[str, Any],
        match_result: Dict[str, Any],
    ) -> Dict[str, Any]:
        """Generate basic optimization without AI (no credit cost)."""
        original = resume.parsed_data or {}
        optimized = original.copy()
        
        # Just highlight matching skills
        optimized["_optimization"] = {
            "type": "basic",
            "matched_skills": match_result.get("matched_skills", []),
            "optimized_at": datetime.utcnow().isoformat(),
        }
        
        return optimized
    
    def _get_next_version_number(self, resume_id: int) -> int:
        """Get the next version number for a resume."""
        latest = (
            self.db.query(CustomizedResume)
            .filter(CustomizedResume.resume_id == resume_id)
            .order_by(CustomizedResume.version_number.desc())
            .first()
        )
        
        if latest:
            return latest.version_number + 1
        return 1
    
    def get_customized_versions(
        self,
        user_id: int,
        resume_id: Optional[int] = None,
    ) -> List[Dict[str, Any]]:
        """Get all customized versions for a user."""
        query = self.db.query(CustomizedResume).filter(
            CustomizedResume.user_id == user_id
        )
        
        if resume_id:
            query = query.filter(CustomizedResume.resume_id == resume_id)
        
        versions = query.order_by(CustomizedResume.created_at.desc()).all()
        
        return [
            {
                "id": v.id,
                "resume_id": v.resume_id,
                "application_id": v.application_id,
                "version_number": v.version_number,
                "target_job_title": v.target_job_title,
                "target_company": v.target_company,
                "created_at": v.created_at.isoformat(),
                "changes_count": len(v.changes_made) if v.changes_made else 0,
            }
            for v in versions
        ]
    
    def generate_pdf(
        self,
        customized_resume_id: int,
        user_id: int,
    ) -> Tuple[Optional[bytes], Optional[str]]:
        """
        Generate PDF from customized resume.
        
        Returns:
            Tuple of (PDF bytes, error message)
        """
        if not FITZ_AVAILABLE:
            return None, "PDF generation not available (PyMuPDF not installed)"
        
        # Get customized resume
        customized = (
            self.db.query(CustomizedResume)
            .filter(
                CustomizedResume.id == customized_resume_id,
                CustomizedResume.user_id == user_id,
            )
            .first()
        )
        
        if not customized:
            return None, "Customized resume not found"
        
        content = customized.customized_content or customized.original_content or {}
        
        try:
            pdf_bytes = self._create_pdf(content, customized)
            return pdf_bytes, None
        except Exception as e:
            logger.error(f"PDF generation error: {e}")
            return None, f"PDF generation failed: {str(e)}"
    
    def _create_pdf(
        self,
        content: Dict[str, Any],
        customized: CustomizedResume,
    ) -> bytes:
        """Create PDF using PyMuPDF."""
        doc = fitz.open()
        
        # Page setup
        page = doc.new_page(width=612, height=792)  # Letter size
        
        # Colors
        header_color = (0.1, 0.1, 0.4)  # Dark blue
        text_color = (0, 0, 0)
        accent_color = (0.2, 0.4, 0.6)
        
        y_pos = 50
        left_margin = 50
        right_margin = 562
        
        # Helper function to insert text
        def insert_text(text, x, y, fontsize=10, color=text_color, fontname="helv"):
            page.insert_text(
                (x, y),
                text,
                fontsize=fontsize,
                color=color,
                fontname=fontname,
            )
            return y + fontsize + 4
        
        # Helper to draw line
        def draw_line(y):
            page.draw_line(
                (left_margin, y),
                (right_margin, y),
                color=accent_color,
                width=0.5,
            )
            return y + 10
        
        # Name (large header)
        name = content.get("name", content.get("full_name", "Candidate"))
        y_pos = insert_text(name, left_margin, y_pos, fontsize=20, color=header_color)
        
        # Contact info
        contact_parts = []
        if content.get("email"):
            contact_parts.append(content["email"])
        if content.get("phone"):
            contact_parts.append(content["phone"])
        if content.get("location"):
            contact_parts.append(content["location"])
        
        if contact_parts:
            y_pos = insert_text(" | ".join(contact_parts), left_margin, y_pos + 5, fontsize=9)
        
        y_pos = draw_line(y_pos + 5)
        
        # Target job title (if optimized)
        if customized.target_job_title:
            y_pos = insert_text(
                f"Applying for: {customized.target_job_title}",
                left_margin, y_pos,
                fontsize=11,
                color=accent_color,
            )
            if customized.target_company:
                y_pos = insert_text(
                    f"at {customized.target_company}",
                    left_margin + 100, y_pos - 12,
                    fontsize=11,
                    color=accent_color,
                )
            y_pos += 10
        
        # Summary section
        if content.get("summary"):
            y_pos = insert_text("SUMMARY", left_margin, y_pos, fontsize=12, color=header_color)
            y_pos = draw_line(y_pos)
            
            summary = content["summary"]
            # Wrap text
            words = summary.split()
            line = ""
            for word in words:
                if len(line + " " + word) < 80:
                    line += " " + word if line else word
                else:
                    y_pos = insert_text(line, left_margin, y_pos, fontsize=10)
                    line = word
            if line:
                y_pos = insert_text(line, left_margin, y_pos, fontsize=10)
            y_pos += 10
        
        # Skills section
        skills = content.get("skills", [])
        if skills:
            y_pos = insert_text("SKILLS", left_margin, y_pos, fontsize=12, color=header_color)
            y_pos = draw_line(y_pos)
            
            skills_text = ", ".join(skills[:15])  # Limit to 15 skills
            # Wrap if needed
            if len(skills_text) > 80:
                y_pos = insert_text(skills_text[:80], left_margin, y_pos, fontsize=10)
                y_pos = insert_text(skills_text[80:160], left_margin, y_pos, fontsize=10)
            else:
                y_pos = insert_text(skills_text, left_margin, y_pos, fontsize=10)
            y_pos += 15
        
        # Experience section
        experience = content.get("experience", [])
        if experience:
            y_pos = insert_text("EXPERIENCE", left_margin, y_pos, fontsize=12, color=header_color)
            y_pos = draw_line(y_pos)
            
            for exp in experience[:4]:  # Limit to 4 experiences
                if y_pos > 700:  # New page if needed
                    page = doc.new_page(width=612, height=792)
                    y_pos = 50
                
                title = exp.get("title", exp.get("position", ""))
                company = exp.get("company", "")
                dates = exp.get("dates", exp.get("duration", ""))
                
                y_pos = insert_text(f"{title} at {company}", left_margin, y_pos, fontsize=11, color=header_color)
                y_pos = insert_text(dates, left_margin, y_pos, fontsize=9, color=accent_color)
                
                # Bullet points
                bullets = exp.get("bullets", exp.get("responsibilities", []))
                for bullet in bullets[:4]:
                    if y_pos > 750:
                        break
                    y_pos = insert_text(f"• {bullet[:100]}", left_margin + 10, y_pos, fontsize=9)
                
                y_pos += 10
        
        # Education section
        education = content.get("education", [])
        if education:
            if y_pos > 700:
                page = doc.new_page(width=612, height=792)
                y_pos = 50
            
            y_pos = insert_text("EDUCATION", left_margin, y_pos, fontsize=12, color=header_color)
            y_pos = draw_line(y_pos)
            
            for edu in education[:3]:
                degree = edu.get("degree", "")
                school = edu.get("school", edu.get("institution", ""))
                year = edu.get("year", edu.get("graduation_year", ""))
                
                y_pos = insert_text(f"{degree}", left_margin, y_pos, fontsize=11)
                y_pos = insert_text(f"{school} - {year}", left_margin, y_pos, fontsize=9, color=accent_color)
                y_pos += 5
        
        # Footer
        page.insert_text(
            (left_margin, 770),
            f"Customized for {customized.target_company or 'target company'} • Version {customized.version_number}",
            fontsize=8,
            color=accent_color,
        )
        
        # Save to bytes
        pdf_bytes = doc.tobytes()
        doc.close()
        
        return pdf_bytes
    
    def generate_docx(
        self,
        customized_resume_id: int,
        user_id: int,
    ) -> Tuple[Optional[bytes], Optional[str]]:
        """
        Generate DOCX from customized resume.
        
        Returns:
            Tuple of (DOCX bytes, error message)
        """
        if not DOCX_AVAILABLE:
            return None, "DOCX generation not available (python-docx not installed)"
        
        # Get customized resume
        customized = (
            self.db.query(CustomizedResume)
            .filter(
                CustomizedResume.id == customized_resume_id,
                CustomizedResume.user_id == user_id,
            )
            .first()
        )
        
        if not customized:
            return None, "Customized resume not found"
        
        content = customized.customized_content or customized.original_content or {}
        
        try:
            docx_bytes = self._create_docx(content, customized)
            return docx_bytes, None
        except Exception as e:
            logger.error(f"DOCX generation error: {e}")
            return None, f"DOCX generation failed: {str(e)}"
    
    def _create_docx(
        self,
        content: Dict[str, Any],
        customized: CustomizedResume,
    ) -> bytes:
        """Create DOCX using python-docx."""
        doc = Document()
        
        # Add name as title
        name = content.get("name", content.get("full_name", "Candidate"))
        title = doc.add_heading(name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_parts = []
        if content.get("email"):
            contact_parts.append(content["email"])
        if content.get("phone"):
            contact_parts.append(content["phone"])
        if content.get("location"):
            contact_parts.append(content["location"])
        
        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Target job
        if customized.target_job_title:
            target = doc.add_paragraph()
            target.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = target.add_run(f"Applying for: {customized.target_job_title}")
            run.italic = True
            if customized.target_company:
                run = target.add_run(f" at {customized.target_company}")
                run.italic = True
        
        # Summary
        if content.get("summary"):
            doc.add_heading("Summary", 1)
            doc.add_paragraph(content["summary"])
        
        # Skills
        skills = content.get("skills", [])
        if skills:
            doc.add_heading("Skills", 1)
            doc.add_paragraph(", ".join(skills[:15]))
        
        # Experience
        experience = content.get("experience", [])
        if experience:
            doc.add_heading("Experience", 1)
            
            for exp in experience[:5]:
                title = exp.get("title", exp.get("position", ""))
                company = exp.get("company", "")
                dates = exp.get("dates", exp.get("duration", ""))
                
                exp_para = doc.add_paragraph()
                run = exp_para.add_run(f"{title}")
                run.bold = True
                exp_para.add_run(f" at {company}")
                
                doc.add_paragraph(dates)
                
                bullets = exp.get("bullets", exp.get("responsibilities", []))
                for bullet in bullets[:4]:
                    doc.add_paragraph(bullet, style="List Bullet")
        
        # Education
        education = content.get("education", [])
        if education:
            doc.add_heading("Education", 1)
            
            for edu in education[:3]:
                degree = edu.get("degree", "")
                school = edu.get("school", edu.get("institution", ""))
                year = edu.get("year", edu.get("graduation_year", ""))
                
                edu_para = doc.add_paragraph()
                run = edu_para.add_run(degree)
                run.bold = True
                edu_para.add_run(f"\n{school} - {year}")
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.read()


# Singleton accessor
def get_resume_customization_service(db: Session) -> ResumeCustomizationService:
    """Get resume customization service instance."""
    return ResumeCustomizationService(db)


__all__ = [
    "ResumeCustomizationService",
    "get_resume_customization_service",
]
