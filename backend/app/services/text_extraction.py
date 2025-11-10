"""Text extraction service for parsing PDF, DOCX, and text files"""

import logging
from typing import Optional
import os
from PyPDF2 import PdfReader
from pdfminer.high_level import extract_text as pdfminer_extract
from docx import Document
import re

logger = logging.getLogger(__name__)


class TextExtractionService:
    """Service for extracting text from various file formats"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file using multiple methods"""
        text = ""
        
        try:
            # Try PyPDF2 first (faster)
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            # If PyPDF2 didn't extract much, try pdfminer (more robust)
            if len(text.strip()) < 100:
                text = pdfminer_extract(file_path)
            
            logger.info(f"Extracted {len(text)} characters from PDF")
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            extracted_text = "\n".join(text)
            logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
            return extracted_text.strip()
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_from_txt(file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            logger.info(f"Extracted {len(text)} characters from TXT")
            return text.strip()
            
        except Exception as e:
            logger.error(f"TXT extraction error: {str(e)}")
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    
    @classmethod
    def extract_text(cls, file_path: str, file_type: str) -> str:
        """Extract text based on file type"""
        file_type = file_type.lower()
        
        if file_type == 'pdf':
            return cls.extract_from_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return cls.extract_from_docx(file_path)
        elif file_type == 'txt':
            return cls.extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\-\:\;\(\)\[\]\/\@\#\$\%\&\*\+\=]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
